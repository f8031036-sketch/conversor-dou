#!/usr/bin/env python3
"""
conversor.py — Converte PDF do SEI para DOCX no padrão DOU
Imprensa Nacional: Calibri 9pt, espaço simples
Suporta: meia página (12cm) ou página inteira (25cm)
Autofit de colunas baseado no conteúdo real
"""

import sys, re
from pathlib import Path
from collections import defaultdict

import pdfplumber
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Calibri"
FONT_SIZE = Pt(9)
# Aprox. largura de 1 caractere Calibri 9pt em cm
CHAR_W_CM = 0.21

FOOTER_RE = re.compile(
    r'(^Minuta de Portaria\s|^Este conteúdo não substitui|^Referência: Processo)',
    re.IGNORECASE
)

# Fronteiras das colunas da tabela SSC/MGI (medidas do PDF)
COL_BOUNDS = [0, 247, 354, 405, 461, 9999]
HEADER_ROW = ['UNIDADE/ÓRGÃO', 'CARGO/FUNÇÃO', 'NÍVEL', 'Nº DOC SEI', 'Nº PROCESSO SEI']

# ── Formatação DOCX ───────────────────────────────────────────────────────────

def set_run(run, bold=False):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = bool(bold)
    rpr = run._r.get_or_add_rPr()
    old = rpr.find(qn('w:rFonts'))
    if old is not None: rpr.remove(old)
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), FONT_NAME)
    rpr.insert(0, rf)


def set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, fi=None, li=None):
    pf = p.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = Pt(10)
    if fi is not None: pf.first_line_indent = Cm(fi)
    if li is not None: pf.left_indent       = Cm(li)


def add_p(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, fi=None, li=None):
    p = doc.add_paragraph()
    set_para(p, align=align, fi=fi, li=li)
    if text:
        r = p.add_run(text)
        set_run(r, bold=bool(bold))
    return p


def set_cell_w(cell, cm):
    tc  = cell._tc
    tcp = tc.get_or_add_tcPr()
    w   = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(Cm(cm).pt * 20)))
    w.set(qn('w:type'), 'dxa')
    old = tcp.find(qn('w:tcW'))
    if old is not None: tcp.remove(old)
    tcp.append(w)


def calc_col_widths(rows_data, total_cm, n_cols):
    """
    Calcula larguras das colunas por autofit:
    - Mede o texto mais largo de cada coluna
    - Distribui proporcionalmente dentro do total disponível
    - Aplica mínimo de 1.8cm por coluna
    - Máximo de 50% do total para qualquer coluna
    """
    # Mede o conteúdo máximo de cada coluna
    max_chars = [0] * n_cols
    for row in rows_data:
        for ci, cell in enumerate(row[:n_cols]):
            longest = max((len(line) for line in str(cell or '').split('\n')), default=0)
            if longest > max_chars[ci]:
                max_chars[ci] = longest

    # Garante mínimo de 4 chars por coluna
    max_chars = [max(c, 4) for c in max_chars]

    total_chars = sum(max_chars)
    # Mínimo por posição: cols de dados fixos (SEI, nível) precisam de mínimo maior
    # Detecta se é a tabela CCE/FCE de 5 colunas pelo cabeçalho
    is_ccefce = (n_cols == 5 and rows_data and
                 len(rows_data[0]) >= 5 and
                 'PROCESSO' in str(rows_data[0][4]).upper())
    if is_ccefce:
        # Mínimos específicos: [unidade, cargo, nível, doc_sei, processo_sei]
        col_mins = [3.0, 3.0, 1.8, 1.8, 3.8]
    else:
        col_mins = [1.8] * n_cols
    max_cm = total_cm * 0.50

    # Distribuição proporcional com limites
    widths = [max(col_mins[i] if i < len(col_mins) else 1.8,
               min(max_cm, (max_chars[i] / total_chars) * total_cm))
              for i in range(n_cols)]

    # Ajusta para somar exatamente total_cm — distribui na coluna mais larga
    diff = total_cm - sum(widths)
    idx_max = widths.index(max(widths))
    widths[idx_max] = round(widths[idx_max] + diff, 2)

    return widths


def add_table(doc, rows_data, total_cm=25):
    if not rows_data:
        return None

    n_cols = max(len(r) for r in rows_data)
    # Normaliza todas as linhas para n_cols
    rows_norm = [list(r) + [''] * (n_cols - len(r)) for r in rows_data]

    widths = calc_col_widths(rows_norm, total_cm, n_cols)

    tbl = doc.add_table(rows=0, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Largura total via XML
    tp = tbl._tbl.find(qn('w:tblPr'))
    if tp is None:
        tp = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tp)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), str(int(Cm(total_cm).pt * 20)))
    tw.set(qn('w:type'), 'dxa')
    old = tp.find(qn('w:tblW'))
    if old is not None: tp.remove(old)
    tp.append(tw)

    for row_cells in rows_norm:
        row    = tbl.add_row()
        is_hdr = (row_cells[:len(HEADER_ROW)] == HEADER_ROW)
        for ci, ct in enumerate(row_cells):
            cell = row.cells[ci]
            set_cell_w(cell, widths[ci])
            p = cell.paragraphs[0]
            set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(str(ct or ''))
            set_run(r, bold=bool(is_hdr))
    return tbl


# ── Extração de PDF ───────────────────────────────────────────────────────────

def words_to_lines(words):
    bkt = defaultdict(list)
    for w in words:
        bkt[round(w['top'] / 4) * 4].append(w)
    lines = []
    for key in sorted(bkt):
        ws = sorted(bkt[key], key=lambda x: x['x0'])
        lines.append({
            'top': key, 'words': ws,
            'text': ' '.join(w['text'] for w in ws),
            'x0_min': min(w['x0'] for w in ws),
        })
    return lines


def col_of(x):
    for i in range(len(COL_BOUNDS)-1):
        if COL_BOUNDS[i] <= x < COL_BOUNDS[i+1]: return i
    return len(COL_BOUNDS)-2


def extract_cells_by_bounds(line):
    """Extrai células usando as fronteiras fixas do PDF SSC/MGI."""
    cols = [''] * (len(COL_BOUNDS)-1)
    for w in line['words']:
        ci = col_of(w['x0'])
        cols[ci] = (cols[ci]+' '+w['text']).strip()
    return cols


def extract_cells_generic(line, col_positions):
    """Extrai células usando posições detectadas automaticamente."""
    if not col_positions:
        return [line['text']]
    cols = [''] * len(col_positions)
    for w in line['words']:
        # Encontra coluna mais próxima
        ci = min(range(len(col_positions)),
                 key=lambda i: abs(w['x0'] - col_positions[i]))
        # Verifica se está à direita da posição da coluna
        for i in range(len(col_positions)-1, -1, -1):
            if w['x0'] >= col_positions[i] - 5:
                ci = i; break
        cols[ci] = (cols[ci]+' '+w['text']).strip()
    return cols


def has_data_cols(line):
    """True se a linha tem padrão de dados de tabela CCE/FCE."""
    t = line['text']
    if 'CARGO/FUNÇÃO' in t and 'NÍVEL' in t: return True
    if re.search(r'\b[CF]CE\s+\d+\.\d+', t):  return True
    if re.search(r'\d{5,}/\d{4}-\d{2}', t):   return True
    for w in line['words']:
        if w['x0'] >= 405 and re.match(r'^\d{7,9}$', w['text']): return True
    return False


def is_table_line(line):
    return has_data_cols(line)


def classify_pre_table(line):
    t, x = line['text'].strip(), line['x0_min']
    if re.match(r'^MINISTÉRIO|^Secretaria de Serviços', t) and x > 100: return 'org'
    if re.match(r'^PORTARIA\s+\S+\s+Nº', t, re.I): return 'title'
    if re.match(r'^ANEXO\b', t, re.I) and x > 200: return 'section_label'
    if re.match(r'^(CAPÍTULO|SEÇÃO|SUBSEÇÃO)\s', t, re.I): return 'chapter'
    if re.match(r'^Art\.\s+\d+', t): return 'article'
    if re.match(r'^(Parágrafo único|§\s*\d+)', t): return 'paragrafo'
    if re.match(r'^[IVXivx]+\s*[-–]\s+', t): return 'inciso'
    if re.match(r'^[a-z]\)\s+', t): return 'alinea'
    if re.match(r'^[A-ZÁÉÍÓÚÃÂÊÎÔÛÇÀ\s]+$', t) and 4 < len(t) < 60 and x > 200:
        return 'signature'
    if x > 260: return 'ementa'
    return 'text'


def extract_pages(pdf_path):
    all_lines = []
    in_table  = False
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(
                x_tolerance=3, y_tolerance=3,
                keep_blank_chars=False, use_text_flow=False)
            for line in words_to_lines(words):
                t = line['text'].strip()
                if not t or FOOTER_RE.match(t):
                    line['kind'] = 'skip'; all_lines.append(line); continue
                if not in_table and 'CARGO/FUNÇÃO' in t and 'NÍVEL' in t:
                    in_table = True
                line['kind'] = 'table_row' if in_table else classify_pre_table(line)
                all_lines.append(line)
    return all_lines


# ── Agrupamento ───────────────────────────────────────────────────────────────

def merge_pre_table(lines):
    result, i = [], 0
    while i < len(lines):
        line = lines[i]; k = line['kind']
        if k == 'skip': i += 1; continue

        if k == 'ementa':
            m = line['text']; j = i+1
            while j < len(lines) and lines[j]['kind'] == 'ementa':
                m += ' ' + lines[j]['text'].strip(); j += 1
            result.append({'kind':'ementa','text':m.strip()}); i=j; continue

        if k == 'text':
            m = line['text']; j = i+1
            while j < len(lines):
                nk,nt = lines[j]['kind'], lines[j]['text'].strip()
                if nk=='skip': j+=1; continue
                if nk=='text':
                    ends = m.rstrip()[-1:] in '.;'
                    if not ends or (nt and nt[0].islower()):
                        m += ' '+nt; j+=1; continue
                break
            result.append({'kind':'text','text':m.strip()}); i=j; continue

        if k in ('article','paragrafo'):
            m = line['text']; j = i+1
            while j < len(lines):
                nk,nt = lines[j]['kind'], lines[j]['text'].strip()
                if nk=='skip': j+=1; continue
                if nk=='text': m += ' '+nt; j+=1; continue
                break
            result.append({'kind':k,'text':m.strip()}); i=j; continue

        result.append({'kind':k,'text':line['text'],'line':line}); i+=1
    return result


def build_table_from_rows(table_lines):
    rows        = []
    pending_c0  = ''
    pending_c1  = ''

    for line in table_lines:
        t     = line['text'].strip()
        cells = extract_cells_by_bounds(line)
        c0,c1,c2,c3,c4 = cells
        has_data = bool(c2 or c3 or c4)

        if 'CARGO/FUNÇÃO' in t and 'NÍVEL' in t:
            rows.append(list(HEADER_ROW)); pending_c0=pending_c1=''; continue

        if not has_data:
            if c1 and not c0:
                if rows and rows[-1][1] and rows[-1] != list(HEADER_ROW):
                    rows[-1][1] = (rows[-1][1]+' '+c1).strip()
                else:
                    pending_c1 = (pending_c1+' '+c1).strip()
            elif c0 and not c1:
                pending_c0 = (pending_c0+' '+c0).strip()
            elif c0 and c1:
                pending_c0 = (pending_c0+' '+c0).strip()
                pending_c1 = (pending_c1+' '+c1).strip()
            continue

        final_c0 = (pending_c0+' '+c0).strip() if c0 else pending_c0
        final_c1 = (pending_c1+' '+c1).strip() if c1 else pending_c1
        pending_c0 = pending_c1 = ''
        rows.append([final_c0, final_c1, c2, c3, c4])
    return rows


# ── Construtor DOCX ───────────────────────────────────────────────────────────

def build_docx(pre_items, table_rows, out, formato='inteira'):
    """
    formato: 'inteira' (27×35cm, margens 1cm, tabela 25cm)
             'meia'    (14×35cm, margens 1cm, tabela 12cm)
    """
    doc = Document()
    sec = doc.sections[0]

    if formato == 'meia':
        sec.page_width    = Cm(14)
        sec.page_height   = Cm(35)
        sec.left_margin   = Cm(1)
        sec.right_margin  = Cm(1)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        tbl_cm = 12
    else:
        sec.page_width    = Cm(27)
        sec.page_height   = Cm(35)
        sec.left_margin   = Cm(1)
        sec.right_margin  = Cm(1)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        tbl_cm = 25

    for p in sec.header.paragraphs: p.clear()
    for p in sec.footer.paragraphs: p.clear()

    sty = doc.styles['Normal']
    sty.font.name = FONT_NAME; sty.font.size = FONT_SIZE
    pf = sty.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = Pt(10)

    C, J = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY
    prev = None

    for item in pre_items:
        k, t = item['kind'], item.get('text','').strip()
        if   k == 'org':           add_p(doc, t, bold=True,  align=C)
        elif k == 'title':
            if prev not in (None,'org'): add_p(doc,'')
            add_p(doc, t, bold=True, align=C)
        elif k == 'ementa':        add_p(doc, t, align=J, li=6.0 if formato=='inteira' else 3.0); add_p(doc,'')
        elif k == 'text':          add_p(doc, t, align=J)
        elif k == 'signature':     add_p(doc,''); add_p(doc, t, bold=True, align=C)
        elif k == 'section_label': add_p(doc,''); add_p(doc, t, bold=True, align=C)
        elif k == 'chapter':       add_p(doc, t, align=C)
        elif k == 'article':       add_p(doc, t, align=J, fi=1.25)
        elif k == 'paragrafo':     add_p(doc, t, align=J, fi=1.25)
        elif k == 'inciso':        add_p(doc, t, align=J, li=1.25)
        elif k == 'alinea':        add_p(doc, t, align=J, li=2.5)
        prev = k

    if table_rows:
        add_p(doc,'')
        add_table(doc, table_rows, total_cm=tbl_cm)
        add_p(doc,'')

    doc.save(out)
    print(f'Salvo: {out} (formato={formato}, tabela={tbl_cm}cm)')


def convert(pdf_path, docx_path=None, formato='inteira'):
    pdf_path  = Path(pdf_path)
    docx_path = Path(docx_path or (pdf_path.stem+'_DOU.docx'))
    print(f'Lendo {pdf_path.name}...')
    lines = extract_pages(pdf_path)
    pre_lines   = [l for l in lines if l['kind'] != 'table_row']
    table_lines = [l for l in lines if l['kind'] == 'table_row']
    print(f'  {len(pre_lines)} linhas texto, {len(table_lines)} linhas tabela')
    pre_items  = merge_pre_table(pre_lines)
    table_rows = build_table_from_rows(table_lines) if table_lines else []
    print(f'  {len(pre_items)} blocos, {len(table_rows)} linhas de tabela')
    build_docx(pre_items, table_rows, str(docx_path), formato=formato)
    return str(docx_path)


if __name__ == '__main__':
    fmt = sys.argv[3] if len(sys.argv) > 3 else 'inteira'
    if len(sys.argv) < 2:
        print('Uso: python3 conversor.py arquivo.pdf [saida.docx] [inteira|meia]')
        sys.exit(1)
    print('Concluído:', convert(sys.argv[1],
                                sys.argv[2] if len(sys.argv)>2 else None,
                                formato=fmt))
