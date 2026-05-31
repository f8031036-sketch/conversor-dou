#!/usr/bin/env python3
"""
conversor_docx.py — Converte DOCX (gerado pelo Word a partir do PDF do SEI)
para DOCX formatado no padrão DOU.
Entrada: .docx salvo pelo Word
Saída:   .docx com Calibri 9pt, espaço simples, tabela única, página 27×35cm
"""

import sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Calibri"
FONT_SIZE = Pt(9)
HEADER_COLS = ['UNIDADE/ÓRGÃO', 'CARGO/FUNÇÃO', 'NÍVEL', 'Nº DOC SEI', 'Nº PROCESSO SEI']

# ── Formatação DOCX ───────────────────────────────────────────────────────────

def set_run(run, bold=False):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = bool(bold)
    rpr = run._r.get_or_add_rPr()
    old = rpr.find(qn('w:rFonts'))
    if old is not None: rpr.remove(old)
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), FONT_NAME)
    rpr.insert(0, rf)


def set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, fi=None, li=None):
    pf = p.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = Pt(10)
    if fi is not None: pf.first_line_indent = Cm(fi)
    if li is not None: pf.left_indent       = Cm(li)


def add_p(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, fi=None, li=None):
    p = doc.add_paragraph()
    set_para(p, align=align, fi=fi, li=li)
    if text:
        r = p.add_run(text)
        set_run(r, bold=bool(bold))
    return p


def set_cell_w(cell, cm):
    tc  = cell._tc
    tcp = tc.get_or_add_tcPr()
    w   = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(Cm(cm).pt * 20)))
    w.set(qn('w:type'), 'dxa')
    old = tcp.find(qn('w:tcW'))
    if old is not None: tcp.remove(old)
    tcp.append(w)


def calc_col_widths(rows_data, total_cm, n_cols):
    """Autofit: distribui proporcionalmente ao conteúdo real."""
    max_chars = [0] * n_cols
    for row in rows_data:
        for ci, cell in enumerate(row[:n_cols]):
            longest = max((len(line) for line in str(cell or '').split('\n')), default=0)
            if longest > max_chars[ci]:
                max_chars[ci] = longest
    max_chars = [max(c, 4) for c in max_chars]
    total_chars = sum(max_chars)

    # Mínimos por tipo de coluna (detecta tabela CCE/FCE de 5 colunas)
    is_ccefce = (n_cols == 5 and rows_data and
                 len(rows_data[0]) >= 5 and
                 'PROCESSO' in str(rows_data[0][4]).upper())
    col_mins = [3.0, 3.0, 1.8, 1.8, 3.8] if is_ccefce else [1.8] * n_cols
    max_pct  = total_cm * 0.50

    widths = [max(col_mins[i] if i < len(col_mins) else 1.8,
                  min(max_pct, (max_chars[i] / total_chars) * total_cm))
              for i in range(n_cols)]

    # Ajusta soma para total_cm
    diff = total_cm - sum(widths)
    widths[widths.index(max(widths))] += round(diff, 2)
    return widths


def add_table_to_doc(doc, rows_data, total_cm=25):
    if not rows_data: return None
    n_cols    = max(len(r) for r in rows_data)
    rows_norm = [list(r) + [''] * (n_cols - len(r)) for r in rows_data]
    widths    = calc_col_widths(rows_norm, total_cm, n_cols)

    tbl = doc.add_table(rows=0, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    tp = tbl._tbl.find(qn('w:tblPr'))
    if tp is None: tp = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tp)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), str(int(Cm(total_cm).pt * 20)))
    tw.set(qn('w:type'), 'dxa')
    old = tp.find(qn('w:tblW'))
    if old is not None: tp.remove(old)
    tp.append(tw)

    for row_cells in rows_norm:
        row    = tbl.add_row()
        is_hdr = (row_cells[:len(HEADER_COLS)] == HEADER_COLS)
        for ci, ct in enumerate(row_cells):
            cell = row.cells[ci]
            set_cell_w(cell, widths[ci])
            p = cell.paragraphs[0]
            set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT)
            # Limpa quebras de linha internas (Word coloca \n em células multiline)
            text = str(ct or '').replace('\n', ' ').strip()
            r = p.add_run(text)
            set_run(r, bold=bool(is_hdr))
    return tbl


# ── Extração do DOCX do Word ──────────────────────────────────────────────────

def classify_para(p):
    """Classifica parágrafo do DOCX gerado pelo Word."""
    t = p.text.strip()
    if not t: return 'vazio'
    if re.match(r'^MINISTÉRIO|^Secretaria de Serviços', t): return 'org'
    if re.match(r'^PORTARIA\s+\S+\s+Nº', t, re.I): return 'title'
    if re.match(r'^ANEXO\b', t, re.I): return 'section_label'
    if re.match(r'^(CAPÍTULO|SEÇÃO|SUBSEÇÃO)\s', t, re.I): return 'chapter'
    if re.match(r'^Art\.\s+\d+', t): return 'article'
    if re.match(r'^(Parágrafo único|§\s*\d+)', t): return 'paragrafo'
    if re.match(r'^[IVXivx]+\s*[-–]\s+', t): return 'inciso'
    if re.match(r'^[a-z]\)\s+', t): return 'alinea'
    bold = any(r.bold for r in p.runs if r.text.strip())
    if bold and re.match(r'^[A-ZÁÉÍÓÚÃÂÊÎÔÛÇÀ\s]+$', t) and len(t) < 60:
        return 'signature'
    # Ementa: texto recuado (indentação > 0) ou alinhado à direita
    if p.paragraph_format.left_indent and p.paragraph_format.left_indent > 0:
        return 'ementa'
    if p.alignment and p.alignment.name == 'RIGHT':
        return 'ementa'
    return 'text'


def extract_from_docx(src_path):
    """
    Lê o DOCX gerado pelo Word e retorna:
    - pre_items: lista de blocos de texto classificados
    - all_table_rows: todas as linhas de todas as tabelas unificadas
    """
    doc = Document(src_path)

    # ── Parágrafos de texto ──
    pre_items = []
    ementa_buffer = []
    in_ementa = False

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue

        kind = classify_para(p)

        if kind == 'title':
            # Após o título vem a ementa
            pre_items.append({'kind': kind, 'text': t})
            in_ementa = True
            continue

        if in_ementa and kind not in ('article', 'text', 'signature',
                                       'section_label', 'chapter'):
            # Acumula linhas da ementa (o Word pode quebrar em vários parágrafos)
            ementa_buffer.append(t)
            continue
        else:
            if ementa_buffer:
                pre_items.append({'kind': 'ementa', 'text': ' '.join(ementa_buffer)})
                ementa_buffer = []
                in_ementa = False

        if kind == 'vazio': continue
        pre_items.append({'kind': kind, 'text': t})

    if ementa_buffer:
        pre_items.append({'kind': 'ementa', 'text': ' '.join(ementa_buffer)})

    # ── Tabelas: une todas em uma lista única ──
    all_rows = []
    first_table = True

    for tbl in doc.tables:
        tbl_rows = []
        for row in tbl.rows:
            cells = [c.text.replace('\n', ' ').strip() for c in row.cells]
            # Pula linhas totalmente vazias
            if not any(cells): continue
            tbl_rows.append(cells)

        if not tbl_rows: continue

        # Verifica se é tabela CCE/FCE (tem cabeçalho com CARGO/FUNÇÃO)
        first_cell = tbl_rows[0][0].strip()
        has_header_row = any('CARGO/FUNÇÃO' in str(r) for r in tbl_rows[:2])

        if has_header_row:
            # Normaliza cabeçalho para padrão
            for ri, row in enumerate(tbl_rows):
                if 'CARGO/FUNÇÃO' in str(row):
                    # Substitui pelo cabeçalho padrão
                    n = len(row)
                    tbl_rows[ri] = HEADER_COLS[:n] + [''] * max(0, n - len(HEADER_COLS))
                    break

        # Se não é a primeira tabela, remove o cabeçalho duplicado de colunas
        if not first_table and has_header_row:
            # Mantém só a linha do nome da seção (primeira) e remove cabeçalho de colunas
            new_rows = []
            for row in tbl_rows:
                if row[:len(HEADER_COLS)] == HEADER_COLS:
                    continue  # Remove cabeçalho repetido
                new_rows.append(row)
            tbl_rows = new_rows

        all_rows.extend(tbl_rows)
        first_table = False

    # Adiciona cabeçalho único no início se não existir
    if all_rows and all_rows[0][:len(HEADER_COLS)] != HEADER_COLS:
        all_rows.insert(0, HEADER_COLS)

    return pre_items, all_rows


# ── Construtor DOCX ───────────────────────────────────────────────────────────

def build_docx(pre_items, table_rows, out, formato='inteira'):
    doc = Document()
    sec = doc.sections[0]

    if formato == 'meia':
        sec.page_width    = Cm(14)
        sec.page_height   = Cm(35)
        tbl_cm = 12
        ementa_li = 3.0
    else:
        sec.page_width    = Cm(27)
        sec.page_height   = Cm(35)
        tbl_cm = 25
        ementa_li = 6.0

    sec.left_margin   = Cm(1)
    sec.right_margin  = Cm(1)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    for p in sec.header.paragraphs: p.clear()
    for p in sec.footer.paragraphs: p.clear()

    sty = doc.styles['Normal']
    sty.font.name = FONT_NAME; sty.font.size = FONT_SIZE
    pf = sty.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = Pt(10)

    C, J = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY
    prev = None

    for item in pre_items:
        k, t = item['kind'], item.get('text', '').strip()
        if not t: continue

        if   k == 'org':           add_p(doc, t, bold=True,  align=C)
        elif k == 'title':
            if prev not in (None, 'org'): add_p(doc, '')
            add_p(doc, t, bold=True, align=C)
        elif k == 'ementa':        add_p(doc, t, align=J, li=ementa_li); add_p(doc, '')
        elif k == 'text':          add_p(doc, t, align=J)
        elif k == 'signature':     add_p(doc, ''); add_p(doc, t, bold=True, align=C)
        elif k == 'section_label': add_p(doc, ''); add_p(doc, t, bold=True, align=C)
        elif k == 'chapter':       add_p(doc, t, align=C)
        elif k == 'article':       add_p(doc, t, align=J, fi=1.25)
        elif k == 'paragrafo':     add_p(doc, t, align=J, fi=1.25)
        elif k == 'inciso':        add_p(doc, t, align=J, li=1.25)
        elif k == 'alinea':        add_p(doc, t, align=J, li=2.5)
        prev = k

    if table_rows:
        add_p(doc, '')
        add_table_to_doc(doc, table_rows, total_cm=tbl_cm)
        add_p(doc, '')

    doc.save(out)
    print(f'Salvo: {out} (formato={formato}, tabela={tbl_cm}cm)')


# ── Main ──────────────────────────────────────────────────────────────────────

def convert_docx(src_path, out_path=None, formato='inteira'):
    src_path = Path(src_path)
    out_path = Path(out_path or (src_path.stem + '_DOU.docx'))
    print(f'Lendo {src_path.name}...')
    pre_items, table_rows = extract_from_docx(src_path)
    print(f'  {len(pre_items)} blocos de texto, {len(table_rows)} linhas de tabela')
    build_docx(pre_items, table_rows, str(out_path), formato=formato)
    return str(out_path)


if __name__ == '__main__':
    fmt = sys.argv[3] if len(sys.argv) > 3 else 'inteira'
    if len(sys.argv) < 2:
        print('Uso: python3 conversor_docx.py arquivo.docx [saida.docx] [inteira|meia]')
        sys.exit(1)
    print('Concluído:', convert_docx(sys.argv[1],
                                     sys.argv[2] if len(sys.argv) > 2 else None,
                                     formato=fmt))
